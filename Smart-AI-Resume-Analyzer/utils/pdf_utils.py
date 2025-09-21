"""
PDF Processing Utilities for Streamlit Cloud Compatibility
Provides fallback PDF processing when PyMuPDF is not available
"""

import io
import os
from typing import Optional, Dict, Any, List, Tuple

# Try multiple PDF processing libraries for cloud compatibility
PDF_PROCESSORS = []

# Primary: pdfplumber (most reliable for text extraction)
try:
    import pdfplumber
    PDF_PROCESSORS.append('pdfplumber')
except ImportError:
    pass

# Secondary: pypdf (good for basic text extraction)
try:
    import pypdf
    PDF_PROCESSORS.append('pypdf')
except ImportError:
    pass

# Tertiary: pdfminer (robust for complex PDFs)  
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.layout import LAParams
    PDF_PROCESSORS.append('pdfminer')
except ImportError:
    pass

# Legacy: PyMuPDF/fitz (if available)
try:
    import fitz
    PDF_PROCESSORS.append('fitz')
except ImportError:
    pass


def extract_text_from_pdf(pdf_file) -> str:
    """
    Extract text from PDF using available processors
    Tries multiple libraries for maximum compatibility
    """
    if not PDF_PROCESSORS:
        raise RuntimeError("No PDF processing libraries available")
    
    # Reset file pointer
    if hasattr(pdf_file, 'seek'):
        pdf_file.seek(0)
    
    # Try each processor in order of preference
    for processor in PDF_PROCESSORS:
        try:
            if processor == 'pdfplumber':
                return _extract_with_pdfplumber(pdf_file)
            elif processor == 'pypdf':
                return _extract_with_pypdf(pdf_file)
            elif processor == 'pdfminer':
                return _extract_with_pdfminer(pdf_file)
            elif processor == 'fitz':
                return _extract_with_fitz(pdf_file)
        except Exception as e:
            print(f"❌ Failed to extract with {processor}: {str(e)}")
            continue
    
    raise RuntimeError("All PDF processors failed to extract text")


def _extract_with_pdfplumber(pdf_file) -> str:
    """Extract text using pdfplumber"""
    text = ""
    
    # Handle both file-like objects and bytes
    if hasattr(pdf_file, 'read'):
        pdf_bytes = pdf_file.read()
    else:
        pdf_bytes = pdf_file
    
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    
    return text.strip()


def _extract_with_pypdf(pdf_file) -> str:
    """Extract text using pypdf"""
    text = ""
    
    # Handle both file-like objects and bytes
    if hasattr(pdf_file, 'read'):
        pdf_bytes = pdf_file.read()
    else:
        pdf_bytes = pdf_file
    
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    for page in reader.pages:
        text += page.extract_text() + "\n"
    
    return text.strip()


def _extract_with_pdfminer(pdf_file) -> str:
    """Extract text using pdfminer"""
    # Handle both file-like objects and bytes
    if hasattr(pdf_file, 'read'):
        pdf_bytes = pdf_file.read()
    else:
        pdf_bytes = pdf_file
    
    laparams = LAParams(
        boxes_flow=0.5,
        word_margin=0.1,
        char_margin=2.0,
        line_margin=0.5
    )
    
    return pdfminer_extract_text(io.BytesIO(pdf_bytes), laparams=laparams)


def _extract_with_fitz(pdf_file) -> str:
    """Extract text using PyMuPDF/fitz (if available)"""
    text = ""
    
    # Handle both file-like objects and bytes
    if hasattr(pdf_file, 'read'):
        pdf_bytes = pdf_file.read()
    else:
        pdf_bytes = pdf_file
    
    doc = fitz.open("pdf", pdf_bytes)
    for page in doc:
        text += page.get_text() + "\n"
    doc.close()
    
    return text.strip()


def create_simple_annotated_pdf(original_pdf, annotations: List[Dict]) -> bytes:
    """
    Create annotated PDF with cloud-compatible methods
    Uses reportlab to overlay annotations on original PDF
    """
    try:
        # First try PyMuPDF if available
        if 'fitz' in PDF_PROCESSORS:
            return _create_fitz_annotated_pdf(original_pdf, annotations)
        
        # Cloud-compatible method using reportlab
        print("🔄 Using cloud-compatible PDF annotation with reportlab...")
        return _create_reportlab_annotated_pdf(original_pdf, annotations)
        
    except Exception as e:
        print(f"❌ PDF annotation failed: {str(e)}")
        # Return original PDF if all methods fail
        if hasattr(original_pdf, 'read'):
            original_pdf.seek(0)
            return original_pdf.read()
        return original_pdf


def _create_reportlab_annotated_pdf(original_pdf, annotations: List[Dict]) -> bytes:
    """Create annotated PDF using reportlab - cloud compatible"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.pdfbase.pdfmetrics import stringWidth
        import tempfile
        
        # Get original PDF bytes
        if hasattr(original_pdf, 'read'):
            original_pdf.seek(0)
            pdf_bytes = original_pdf.read()
        else:
            pdf_bytes = original_pdf
        
        # Create a temporary file for the overlay
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as overlay_file:
            overlay_path = overlay_file.name
            
            # Create overlay PDF with annotations
            c = canvas.Canvas(overlay_path, pagesize=letter)
            page_width, page_height = letter
            
            # Group annotations by page
            page_annotations = {}
            for annotation in annotations:
                page_num = annotation.get('page', 0)
                if page_num not in page_annotations:
                    page_annotations[page_num] = []
                page_annotations[page_num].append(annotation)
            
            # Create overlay for each page with annotations
            for page_num, page_annots in page_annotations.items():
                # Add annotations to this page
                for annotation in page_annots:
                    note = annotation.get('note', 'Feedback')
                    color = annotation.get('color', [1, 1, 0])  # Default yellow
                    rect = annotation.get('rect', [50, page_height-100, 400, page_height-50])
                    
                    # Convert color to reportlab color
                    if len(color) >= 3:
                        rl_color = colors.Color(color[0], color[1], color[2], alpha=0.3)
                    else:
                        rl_color = colors.Color(1, 1, 0, alpha=0.3)
                    
                    # Draw highlight rectangle
                    c.setFillColor(rl_color)
                    c.rect(rect[0], rect[1], rect[2]-rect[0], rect[3]-rect[1], fill=1, stroke=0)
                    
                    # Add text annotation
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica", 8)
                    
                    # Word wrap the note text
                    words = note.split()
                    lines = []
                    current_line = ""
                    max_width = rect[2] - rect[0] - 10
                    
                    for word in words:
                        test_line = current_line + " " + word if current_line else word
                        if stringWidth(test_line, "Helvetica", 8) <= max_width:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        lines.append(current_line)
                    
                    # Draw text lines
                    y_offset = rect[3] - 12
                    for line in lines[:3]:  # Limit to 3 lines
                        c.drawString(rect[0] + 5, y_offset, line)
                        y_offset -= 10
                
                # End this page
                c.showPage()
            
            c.save()
        
        # Now merge the overlay with the original PDF
        try:
            import pypdf
            
            # Read original PDF
            original_reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            
            # Read overlay PDF
            with open(overlay_path, 'rb') as overlay_file:
                overlay_reader = pypdf.PdfReader(overlay_file)
                
                # Create output PDF
                output_writer = pypdf.PdfWriter()
                
                # Merge pages
                for page_num, page in enumerate(original_reader.pages):
                    # Get overlay page if it exists
                    if page_num < len(overlay_reader.pages):
                        overlay_page = overlay_reader.pages[page_num]
                        page.merge_page(overlay_page)
                    
                    output_writer.add_page(page)
                
                # Write to bytes
                output_bytes = io.BytesIO()
                output_writer.write(output_bytes)
                annotated_pdf_bytes = output_bytes.getvalue()
        
        except Exception as merge_error:
            print(f"⚠️ PDF merge failed: {merge_error}")
            # Return original PDF if merge fails
            annotated_pdf_bytes = pdf_bytes
        
        # Clean up temporary file
        try:
            os.unlink(overlay_path)
        except:
            pass
        
        print("✅ Cloud-compatible PDF annotation completed")
        return annotated_pdf_bytes
    
    except Exception as e:
        print(f"❌ Reportlab annotation failed: {str(e)}")
        # Return original PDF
        if hasattr(original_pdf, 'read'):
            original_pdf.seek(0)
            return original_pdf.read()
        return original_pdf


def _create_fitz_annotated_pdf(original_pdf, annotations: List[Dict]) -> bytes:
    """Create annotated PDF using PyMuPDF if available"""
    if hasattr(original_pdf, 'read'):
        pdf_bytes = original_pdf.read()
    else:
        pdf_bytes = original_pdf
    
    doc = fitz.open("pdf", pdf_bytes)
    
    for annotation in annotations:
        page_num = annotation.get('page', 0)
        if page_num < len(doc):
            page = doc[page_num]
            
            # Add highlight annotation
            rect = annotation.get('rect', [50, 50, 150, 70])
            highlight = page.add_highlight_annot(fitz.Rect(rect))
            
            # Set color based on feedback type
            color = annotation.get('color', [1, 1, 0])  # Default yellow
            highlight.set_colors(stroke=color)
            
            # Add note
            note = annotation.get('note', 'Feedback')
            highlight.set_content(note)
            highlight.update()
    
    # Save annotated PDF
    annotated_bytes = doc.write()
    doc.close()
    
    return annotated_bytes


def get_available_processors() -> List[str]:
    """Get list of available PDF processors"""
    return PDF_PROCESSORS.copy()


def get_pdf_info() -> Dict[str, Any]:
    """Get information about PDF processing capabilities"""
    return {
        'available_processors': PDF_PROCESSORS,
        'primary_processor': PDF_PROCESSORS[0] if PDF_PROCESSORS else None,
        'annotation_support': 'fitz' in PDF_PROCESSORS,
        'total_processors': len(PDF_PROCESSORS)
    }


# Test function for debugging
def test_pdf_processing():
    """Test PDF processing capabilities"""
    info = get_pdf_info()
    print("📄 PDF Processing Capabilities:")
    print(f"   Available processors: {info['available_processors']}")
    print(f"   Primary processor: {info['primary_processor']}")
    print(f"   Annotation support: {info['annotation_support']}")
    print(f"   Total processors: {info['total_processors']}")
    
    if not info['available_processors']:
        print("❌ No PDF processors available!")
        return False
    
    print("✅ PDF processing ready!")
    return True


def extract_text_from_docx(docx_file) -> str:
    """
    Extract text from DOCX file
    Compatible with Streamlit Cloud environment
    """
    try:
        from docx import Document
        
        # Handle both file-like objects and bytes
        if hasattr(docx_file, 'read'):
            doc = Document(docx_file)
        else:
            doc = Document(io.BytesIO(docx_file))
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    
    except ImportError:
        raise RuntimeError("python-docx package not available for DOCX processing")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")


if __name__ == "__main__":
    test_pdf_processing()